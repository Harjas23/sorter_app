import streamlit as st
from docx import Document
import time
import copy


# Utility function to highlight swapped elements in red
def highlight_swap(arr, indices):
    return [f"**:red[{v}]**" if i in indices else str(v) for i, v in enumerate(arr)]


# Bubble Sort with detailed explanations
def bubble_sort(arr):
    steps = []
    n = len(arr)
    for i in range(n):
        swapped = False
        step_message = f"Pass {i + 1}: Comparing elements to place the {'smallest' if i == 0 else f'{i + 1}th smallest'} element."
        for j in range(0, n - i - 1):
            comparison = f"Comparing {arr[j]} and {arr[j + 1]}."
            if arr[j] > arr[j + 1]:
                arr[j], arr[j + 1] = arr[j + 1], arr[j]
                swapped = True
                steps.append((copy.deepcopy(arr), (j, j + 1), f"{comparison} Swapped."))
            else:
                steps.append((copy.deepcopy(arr), (), f"{comparison} No swap."))
        if not swapped:
            steps.append((copy.deepcopy(arr), (), f"No swaps in this pass. Array is sorted."))
            break
    return steps


# Insertion Sort with detailed explanations
def insertion_sort(arr):
    steps = []
    for i in range(1, len(arr)):
        key = arr[i]
        j = i - 1
        steps.append((copy.deepcopy(arr), (), f"Inserting {key} into sorted part of the array."))
        while j >= 0 and key < arr[j]:
            arr[j + 1] = arr[j]
            steps.append((copy.deepcopy(arr), (j, j + 1), f"Moved {arr[j]} to position {j + 1}."))
            j -= 1
        arr[j + 1] = key
        steps.append((copy.deepcopy(arr), (), f"Placed {key} at position {j + 1}."))
    return steps


# Selection Sort with detailed explanations
def selection_sort(arr):
    steps = []
    n = len(arr)
    for i in range(n):
        min_idx = i
        steps.append((copy.deepcopy(arr), (), f"Pass {i + 1}: Finding the smallest element from index {i} onward."))
        for j in range(i + 1, n):
            if arr[j] < arr[min_idx]:
                min_idx = j
        if min_idx != i:
            arr[i], arr[min_idx] = arr[min_idx], arr[i]
            steps.append((copy.deepcopy(arr), (i, min_idx), f"Swapped {arr[min_idx]} and {arr[i]}."))
    return steps


# Quick Sort with detailed explanations
def quick_sort(arr):
    steps = []

    def _quick_sort(arr, low, high):
        if low < high:
            pivot_index = partition(arr, low, high)
            steps.append((copy.deepcopy(arr), (), f"Pivot placed at index {pivot_index}."))
            _quick_sort(arr, low, pivot_index - 1)
            _quick_sort(arr, pivot_index + 1, high)

    def partition(arr, low, high):
        pivot = arr[high]
        i = low - 1
        steps.append((copy.deepcopy(arr), (), f"Choosing pivot {pivot} at index {high}."))
        for j in range(low, high):
            if arr[j] < pivot:
                i += 1
                arr[i], arr[j] = arr[j], arr[i]
                steps.append((copy.deepcopy(arr), (i, j), f"Swapped {arr[j]} with {arr[i]}."))
        arr[i + 1], arr[high] = arr[high], arr[i + 1]
        return i + 1

    _quick_sort(arr, 0, len(arr) - 1)
    return steps


# Merge Sort with detailed explanations
def merge_sort(arr):
    steps = []

    def _merge_sort(arr, left, right):
        if left < right:
            mid = (left + right) // 2
            steps.append((copy.deepcopy(arr), (), f"Splitting array into two halves: {arr[left:mid + 1]} and {arr[mid + 1:right + 1]}."))
            _merge_sort(arr, left, mid)
            _merge_sort(arr, mid + 1, right)
            merge(arr, left, mid, right)

    def merge(arr, left, mid, right):
        left_half = arr[left:mid + 1]
        right_half = arr[mid + 1:right + 1]
        steps.append((copy.deepcopy(arr), (), f"Merging {left_half} and {right_half}."))
        i = 0
        j = 0
        k = left
        while i < len(left_half) and j < len(right_half):
            if left_half[i] <= right_half[j]:
                arr[k] = left_half[i]
                i += 1
            else:
                arr[k] = right_half[j]
                j += 1
            k += 1
        while i < len(left_half):
            arr[k] = left_half[i]
            i += 1
            k += 1
        while j < len(right_half):
            arr[k] = right_half[j]
            j += 1
            k += 1

    _merge_sort(arr, 0, len(arr) - 1)
    return steps


# Function to generate a Word document with sorting steps
def generate_word_doc(steps, algorithm_name):
    doc = Document()
    doc.add_heading(f"{algorithm_name} - Sorting Steps", level=1)
    for idx, (step, indices, explanation) in enumerate(steps):
        doc.add_heading(f"Step {idx + 1}", level=2)
        doc.add_paragraph(explanation)
        highlighted_step = " ".join(highlight_swap(step, indices))
        doc.add_paragraph(highlighted_step)
    return doc


# Streamlit App
st.title("Sorting Algorithm Visualizer")
st.write("Visualize and understand various sorting algorithms step by step.")

# Sorting algorithms and dropdown
array = [456, 125, 59, 23, 45, 9, 234, 6, 450, 389, 148, 78, 28, 97, 208, 109, 2, 5]
sort_methods = {
    "Bubble Sort": bubble_sort,
    "Insertion Sort": insertion_sort,
    "Selection Sort": selection_sort,
    "Quick Sort": quick_sort,
    "Merge Sort": merge_sort,
}

selected_sort = st.selectbox("Select a Sorting Algorithm", list(sort_methods.keys()))

if st.button("Sort"):
    steps = sort_methods[selected_sort](copy.deepcopy(array))
    for idx, (step, indices, explanation) in enumerate(steps):
        st.write(f"Step {idx + 1}: {explanation}")
        st.write("Array:", " ".join(highlight_swap(step, indices)))
        time.sleep(0.5)

    # Generate Word document
    doc = generate_word_doc(steps, selected_sort)
    file_path = f"{selected_sort}_sorting_steps.docx"
    doc.save(file_path)

    with open(file_path, "rb") as file:
        st.download_button(
            label="Download Sorting Steps as Word File",
            data=file,
            file_name=file_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
